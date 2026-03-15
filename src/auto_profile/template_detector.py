from pathlib import Path
from openpyxl import load_workbook
from docx import Document


def detect_template_structure(template_path: str) -> dict:
    ext = Path(template_path).suffix.lower()

    if ext in [".xlsx", ".xlsm"]:
        return detect_excel_structure(template_path)

    if ext == ".docx":
        return detect_word_structure(template_path)

    raise ValueError(f"暂不支持的模板类型：{ext}")


def detect_excel_structure(template_path: str) -> dict:
    wb = load_workbook(template_path)
    ws = wb.active

    max_row = ws.max_row
    max_col = ws.max_column

    # 先判断是否像竖版：A列字段名，B列待填
    a_values = []
    b_values = []

    for r in range(1, min(max_row, 50) + 1):
        a_values.append(ws.cell(r, 1).value)
        b_values.append(ws.cell(r, 2).value if max_col >= 2 else None)

    a_non_empty = [str(v).strip() for v in a_values if v is not None and str(v).strip()]
    b_empty_ratio = sum(1 for v in b_values if v is None or not str(v).strip()) / max(1, len(b_values))

    if len(a_non_empty) >= 3 and b_empty_ratio >= 0.5:
        field_names = []
        for r in range(1, min(max_row, 100) + 1):
            v = ws.cell(r, 1).value
            if v is not None and str(v).strip():
                field_names.append(str(v).strip())

        return {
            "task_mode": "single_record",
            "template_mode": "vertical",
            "field_names": field_names
        }

    # 再判断是否像横向表格：前若干行中找最像表头的一行
    for r in range(1, min(max_row, 15) + 1):
        row_values = [ws.cell(r, c).value for c in range(1, max_col + 1)]
        row_texts = [str(v).strip() for v in row_values if v is not None and str(v).strip()]

        if len(row_texts) < 2:
            continue

        short_count = sum(1 for x in row_texts if len(x) <= 15)
        if short_count >= 2:
            return {
                "task_mode": "table_records",
                "template_mode": "excel_table",
                "header_row": r,
                "start_row": r + 1,
                "field_names": row_texts
            }

    raise ValueError("无法自动识别 Excel 模板结构")


def detect_word_structure(template_path: str) -> dict:
    doc = Document(template_path)

    if not doc.tables:
        raise ValueError("Word 模板中没有表格，暂不支持自动识别纯正文模板")

    table = doc.tables[0]
    if len(table.rows) < 1:
        raise ValueError("Word 模板表格为空")

    header_fields = []
    for cell in table.rows[0].cells:
        text = cell.text.strip()
        if text:
            header_fields.append(text)

    if len(header_fields) < 2:
        raise ValueError("无法识别 Word 表格表头")

    return {
        "task_mode": "table_records",
        "template_mode": "word_table",
        "table_index": 0,
        "header_row": 0,
        "start_row": 1,
        "field_names": header_fields
    }