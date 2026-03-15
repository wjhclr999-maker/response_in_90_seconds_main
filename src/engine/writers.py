from copy import copy
from openpyxl import load_workbook
from docx import Document

def _normalize_records(records):
    if isinstance(records, dict) and "records" in records:
        records = records["records"]

    if not isinstance(records, list):
        raise ValueError("表格填充要求数据为 list[dict] 或 {'records': [...]} 格式")

    for item in records:
        if not isinstance(item, dict):
            raise ValueError("records 中每一项都必须是 dict")

    return records

def fill_excel_vertical(template_path: str, output_path: str, data: dict):
    wb = load_workbook(template_path)
    ws = wb.active

    for row in range(2, ws.max_row + 1):
        field_name = ws[f"A{row}"].value
        if field_name is None:
            continue
        field_name = str(field_name).strip()
        if field_name in data:
            ws[f"B{row}"] = data[field_name]

    wb.save(output_path)

def fill_excel_table(template_path: str, output_path: str, records, header_row: int = 1, start_row: int = 2):
    records = _normalize_records(records)

    wb = load_workbook(template_path)
    ws = wb.active

    header_map = {}
    for col in range(1, ws.max_column + 1):
        cell_value = ws.cell(row=header_row, column=col).value
        if cell_value is not None:
            header_map[str(cell_value).strip()] = col

    if not header_map:
        raise ValueError("Excel 模板表头为空，无法进行表格填充")

    template_style_row = start_row

    for i, record in enumerate(records):
        target_row = start_row + i

        if target_row > ws.max_row:
            ws.insert_rows(target_row)

            for col in range(1, ws.max_column + 1):
                src = ws.cell(row=template_style_row, column=col)
                dst = ws.cell(row=target_row, column=col)
                if src.has_style:
                    dst._style = copy(src._style)
                dst.font = copy(src.font)
                dst.fill = copy(src.fill)
                dst.border = copy(src.border)
                dst.alignment = copy(src.alignment)
                dst.protection = copy(src.protection)
                dst.number_format = src.number_format

        for field_name, value in record.items():
            if field_name in header_map:
                col = header_map[field_name]
                ws.cell(row=target_row, column=col, value=value)

    wb.save(output_path)


def fill_word_table(template_path: str, output_path: str, records, table_index: int = 0, header_row: int = 0, start_row: int = 1):
    records = _normalize_records(records)

    doc = Document(template_path)

    if not doc.tables:
        raise ValueError("Word 模板中没有找到表格")

    if table_index >= len(doc.tables):
        raise ValueError(f"Word 模板只有 {len(doc.tables)} 个表格，table_index={table_index} 越界")

    table = doc.tables[table_index]

    header_map = {}
    header_cells = table.rows[header_row].cells
    for idx, cell in enumerate(header_cells):
        text = cell.text.strip()
        if text:
            header_map[text] = idx

    if not header_map:
        raise ValueError("Word 模板表头为空，无法进行表格填充")

    required_rows = start_row + len(records)
    while len(table.rows) < required_rows:
        table.add_row()

    for i, record in enumerate(records):
        row_idx = start_row + i
        row_cells = table.rows[row_idx].cells

        for field_name, value in record.items():
            if field_name in header_map:
                col_idx = header_map[field_name]
                row_cells[col_idx].text = "" if value is None else str(value)

    doc.save(output_path)